"""Docx tools.
"""
import docx
from typing import List, Dict
from io import BytesIO
from ltp import StnSplit


class DocxTools(object):
    """The class of docx tools.
    """

    def __init__(self, docx_path: BytesIO):
        self.docx = docx.Document(docx_path)
        self.ltp_stnsplitter = StnSplit()
    
    def _paragraphs(self, skip_para_type: List = []) -> List[Dict]:
        """Split paragraphs for docx.

        Returns: The paragraphs of the splited docx.

            para = {
                "para_idx": 20,
                "para_text": "第五条 乙方权利和义务",
                "sents": [
                    {"sent_idx": 0, "sent_text": "第五条"},
                    {"sent_idx": 1, "sent_text": " "},
                    {"sent_idx": 2, "sent_text": "乙方权利和义务"}
                ],
                "para_type": "Heading"
            }
        """
        paragraphs = []
        for index, para in enumerate(self.docx.paragraphs):
            if para.style.name in skip_para_type:
                continue
            paragraph = {
                "para_idx": index,
                "para_text": para.text,
                "sents": [
                    {"sent_idx": idx, "sent_text": sent} 
                    for idx, sent in enumerate(self.ltp_stnsplitter.split(para.text))
                ],
                "para_type": para.style.name
            }
            paragraphs.append(paragraph)
        return paragraphs

    @property
    def paragraphs(self):
        paras = [para["para_text"] for para in self._paragraphs() ]
        return paras
    
    @property
    def sentences(self):
        sents = []
        for para in self._paragraphs():
            for sent in para["sents"]:
                sents.append(sent["sent_text"])
        return sents
    
    def split_sentences(self, text: str) -> List:
        """Split sentences for text.

        Returns: The sentences of the text.
        """
        if not text:
            return []
        sents = self.ltp_stnsplitter.split(text)
        return sents


if __name__ == "__main__":
    docx_obj = DocxTools(
        docx_path="./data/round1_test_data/风城油田夏子街转油站完善工程.docx",
    )
    paras = docx_obj.paragraphs
    sents = docx_obj.sentences
    pass